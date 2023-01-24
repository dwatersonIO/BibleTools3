import docx
import re

name_chapter = re.findall(r'\w+', "Exodus 01.docx")

''' Breaks current chapter file name into list like this: ['Exodus', '02', 'docx']
    Will then use this list below to identify book name and current chapter number '''

BOOK_NAME= name_chapter[0]
print (f'Book name = {BOOK_NAME}')

CHAPTER_NUMBER = int(name_chapter[1])
print (f'Chapter number = {CHAPTER_NUMBER}')


'''
Gets text from Word doc filename using the docx package
Returns:    chapter_text = string of contents 
            len of chapter_text

Appends each para to a list since this is how to interate over file using the docx package.
But we want a string so ... returns string converted from list and replaces /n with a space

'''
filename = "C:/Coding/BibleTools3/Hebrew Scriptures/02-Exodus/Exodus 01.docx"
doc = docx.Document(filename)
chapter_paragraphs = []

# Append to a list

for index, para in enumerate(doc.paragraphs, 1): 
# for para in doc.paragraphs:
    chapter_paragraphs.append(para.text)
 
# Then convert back to a string
RAW_CHAPTER_TEXT = ' '.join(chapter_paragraphs)

if CHAPTER_NUMBER == 1:
    # Check if chapter 1 and get rid of Bible book name from begining of string
    print ("This is Chapter 1")
    CHAPTER_TEXT = ' '.join(RAW_CHAPTER_TEXT.split()[1:])

else:
    # If not chapter one that the RAW_CHAPTER_TEXT does not have Bible book name so fine as is
    CHAPTER_TEXT = RAW_CHAPTER_TEXT

# print (CHAPTER_TEXT)

'''
**********************************************************************
Get the book name from the CHAPTER_TEXT string
'''
# re.search with group method will find the first match only. Will
# book_name = re.search(r'^.+\s', doc).group()

# print (CHAPTER_TEXT)
test = CHAPTER_TEXT.split(' ')

# re.search returns the colon also so need to get rid of that below
# CHAPTER_NUM = chapter_num_with_colon[:-1]

'''
**********************************************************************
Get the chapter number forom the CHAPTER_TEXT string
'''
# re.search with group method will find the first match only. Will
chapter_num_with_colon = re.search(r'\d+:|$', CHAPTER_TEXT).group()

# re.search returns the colon also so need to get rid of that below
CHAPTER_NUM = chapter_num_with_colon[:-1]
print (f'The chapter number is: ', CHAPTER_NUM)

'''
**********************************************************************
Get the number of verses from the CHAPTER_TEXT string
'''

num_of_verses_with_colon = re.findall(r':\d+', CHAPTER_TEXT)[-1]

# Get rid of colon from match and convert to integer
NUM_OF_VERSES = int(num_of_verses_with_colon[1:])

print (f'The number of verses is: ', NUM_OF_VERSES)

'''
**********************************************************************
Get the verses from the CHAPTER_TEXT string

    Uses re module find the search term and adds the verses with that term
    to a dictionary. The dictionaries are then appended to a list that is returned
                 
    Hint:   \d = any digit
            + = one or more of previous character

    '''    
 
this_chapter_dict = {}
list_match_pairs = []

pattern = re.compile(r'\d+:\d+') 
matches = pattern.finditer(CHAPTER_TEXT)

for match in matches:
    list_match_pairs.append((match.start(), match.end()))

for index, pair in enumerate(list_match_pairs, 0): 
    
    # Check if this is the last chapter/verse pair in the list
    if index == NUM_OF_VERSES - 1:   
        start = 1 + list_match_pairs[index][1] 
        # Add 1 to get rid of initial space after verse number

        end = len(CHAPTER_TEXT) # since the last pair, the end of text will be the end of the chapter
        verse_text = CHAPTER_TEXT[start:end]
    
    else:
        # This runs for all chapter/verse pairs if not last one 
        start = 1 + list_match_pairs[index][1]
        # Add 1 to get rid of initial space after verse number
        end = list_match_pairs[index + 1][0]
        verse_text = CHAPTER_TEXT[start:end]
        

    this_chapter_dict["book_name"] = BOOK_NAME
    this_chapter_dict["chapter_num"] = CHAPTER_NUM
    this_chapter_dict["verse_num"] = index + 1 # verse numbers don't start from zero like indices
    this_chapter_dict["verse_text"] = verse_text

    print (this_chapter_dict)


# '''
# Scafolding for testing the search algorithm
# '''
# chapter = get_doc_text("C:/Coding/BibleTools3/Hebrew Scriptures/02-Exodus/Exodus 01.docx") # get the chapter text as a string 

# result = get_all_verses("C:/Coding/BibleTools3/Hebrew Scriptures/02-Exodus","Exodus")

# print (result)

# print (chapter)